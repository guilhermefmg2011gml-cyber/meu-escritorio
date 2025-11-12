# ===========================================================
# pipeline_juridico.py
# GPT-5 + ChromaDB (RAG) + Tavily + DOCX Formatter
# ===========================================================
import os
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Any, Optional

import chromadb
from chromadb.config import Settings
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# prompts por área + busca jurisprudencial abrangente
from services.prompt_juridico import get_prompt_by_area, buscar_jurisprudencia

# ============== ENV / PATHS ==============
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5-turbo")
OPENAI_EMBEDDINGS = os.getenv("OPENAI_EMBEDDINGS", "text-embedding-3-large")
CHROMA_DIR = os.getenv("CHROMA_DIR", "./storage/chroma")

STORAGE_DIR = Path("./storage").resolve()
DOCS_DIR = STORAGE_DIR / "docs"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

# ============== CLIENTES ==============
client_oa: Optional[OpenAI] = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

# ============== CHROMA INIT ==============
chroma = chromadb.Client(Settings(persist_directory=CHROMA_DIR, anonymized_telemetry=False))


class OpenAIEmbeddingFn:
    """Função de embeddings OpenAI para consultas semânticas no Chroma."""

    def __init__(self, client: OpenAI, model: str):
        self.client = client
        self.model = model

    def __call__(self, texts: List[str]) -> List[List[float]]:
        resp = self.client.embeddings.create(model=self.model, input=texts)
        return [d.embedding for d in resp.data]


default_client = client_oa
embedder = OpenAIEmbeddingFn(default_client, OPENAI_EMBEDDINGS) if default_client else None


def get_collection(name: str):
    # Mesmo sem embedder (modo demo), criamos a collection.
    return chroma.get_or_create_collection(name=name, metadata={"hnsw:space": "cosine"})


COL_PECAS = get_collection("pecas_juridicas")  # peças produzidas
COL_MODEL = get_collection("modelos_escritorio")  # modelos internos
COL_JURIS = get_collection("jurisprudencias")  # ementas/trechos internos
COL_DOCS = get_collection("docs_clientes")  # documentos dos clientes


# ============== TIPOS ==============
@dataclass
class RAGContext:
    tipo_peca: str
    facts: str
    pedidos: str
    documentos: List[str]
    cliente_id: Optional[str]
    processo_id: Optional[str]
    area: Optional[str]
    similares: Dict[str, List[Dict[str, Any]]]
    jurisprudencia: List[Dict[str, str]]


# ============== UPSERT / INGESTÃO ==============
def _upsert(collection, texto: str, meta: Dict[str, Any]):
    _id = meta.get("id") or uuid.uuid4().hex
    collection.add(ids=[_id], documents=[texto], metadatas=[meta])
    return _id


def upsert_peca(texto: str, cliente: Optional[str] = None, processo: Optional[str] = None):
    return _upsert(
        COL_PECAS,
        texto,
        {"tipo": "peca", "cliente": cliente or "", "processo": processo or ""},
    )


def upsert_modelo(texto: str, assunto: Optional[str] = None):
    return _upsert(COL_MODEL, texto, {"tipo": "modelo", "assunto": assunto or ""})


def upsert_juris(texto: str, origem: Optional[str] = None, url: Optional[str] = None):
    return _upsert(
        COL_JURIS,
        texto,
        {"tipo": "juris", "origem": origem or "", "url": url or ""},
    )


def upsert_doc_cliente(
    texto: str,
    cliente: Optional[str] = None,
    processo: Optional[str] = None,
    nome: Optional[str] = None,
):
    return _upsert(
        COL_DOCS,
        texto,
        {
            "tipo": "doc",
            "cliente": cliente or "",
            "processo": processo or "",
            "nome": nome or "",
        },
    )


# ============== RETRIEVE (CHROMA) ==============
def search_similares(query: str, k: int = 4) -> Dict[str, List[Dict[str, Any]]]:
    out: Dict[str, List[Dict[str, Any]]] = {}
    for label, col in (
        ("pecas", COL_PECAS),
        ("modelos", COL_MODEL),
        ("juris", COL_JURIS),
        ("docs", COL_DOCS),
    ):
        try:
            res = col.query(query_texts=[query], n_results=k)
            pack: List[Dict[str, Any]] = []
            ids = res.get("ids", [[]])[0]
            docs = res.get("documents", [[]])[0]
            metas = (res.get("metadatas") or [[]])[0]
            for i in range(len(ids)):
                pack.append(
                    {"id": ids[i], "texto": docs[i], "meta": metas[i] if metas else {}}
                )
            out[label] = pack
        except Exception:
            out[label] = []
    return out


# ============== TAVILY WRAPPER ==============
def buscar_tavily(pergunta: str, area: Optional[str] = None, k: int = 7) -> List[Dict[str, str]]:
    """Proxy para a busca de jurisprudência abrangente (superiores, federais, estaduais, trabalhistas)."""

    return buscar_jurisprudencia(pergunta, area=area, k=k)


# ============== PROMPT BUILDER ==============
def _montar_contexto_rag(
    similares: Dict[str, List[Dict[str, Any]]], jurisprudencia: List[Dict[str, str]]
) -> str:
    blocos: List[str] = []

    if similares.get("modelos"):
        blocos.append(
            "=== MODELOS DO ESCRITÓRIO ===\n"
            + "\n\n".join([f"- {d['texto'][:800]}" for d in similares["modelos"]])
        )
    if similares.get("pecas"):
        blocos.append(
            "=== PEÇAS ANTERIORES ===\n"
            + "\n\n".join([f"- {d['texto'][:800]}" for d in similares["pecas"]])
        )
    if similares.get("docs"):
        blocos.append(
            "=== DOCUMENTOS DO CLIENTE ===\n"
            + "\n\n".join([f"- {d['texto'][:700]}" for d in similares["docs"]])
        )
    if similares.get("juris"):
        blocos.append(
            "=== JURIS (BASE INTERNA) ===\n"
            + "\n\n".join([f"- {d['texto'][:600]}" for d in similares["juris"]])
        )
    if jurisprudencia:
        blocos.append(
            "=== JURISPRUDÊNCIA/TEXTOS EXTERNOS (Tavily) ===\n"
            + "\n".join(
                [
                    f"- {j['titulo']} — {j['url']}\n  {j['snippet']}"
                    for j in jurisprudencia
                ]
            )
        )

    return "\n\n".join(blocos) if blocos else "—"


def montar_messages(ctx: RAGContext) -> List[Dict[str, str]]:
    system_prompt = get_prompt_by_area(ctx.area)
    contexto_ext = _montar_contexto_rag(ctx.similares, ctx.jurisprudencia)

    user_prompt = f"""Tipo de peça: {ctx.tipo_peca}

FATOS (resumo do caso):
{ctx.facts}

PEDIDOS:
{ctx.pedidos or 'a) Citação; b) Procedência; c) Custas e honorários.'}

DOCUMENTOS: {", ".join(ctx.documentos) if ctx.documentos else "—"}
CLIENTE: {ctx.cliente_id or "—"} • PROCESSO: {ctx.processo_id or "—"}

CONTEXTOS (RAG):
{contexto_ext}

Instruções:
- Estruture a peça completa (Preâmbulo; I – Dos Fatos; II – Do Direito; III – Dos Pedidos; IV – Provas; V – Valor da Causa, se couber; Termos).
- Faça subsunção norma→fato→conclusão; evite floreios e citações longas literais.
- Use jurisprudência de maneira RESUMIDA quando pertinente.
- Linguagem técnica, objetiva e com urbanidade forense.
"""
    return [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}]


# ============== LLM CALL ==============
def gerar_texto(messages: List[Dict[str, str]]) -> str:
    if not client_oa:
        # Modo DEMO (sem chave): devolve o conteúdo do usuário para depuração.
        for m in messages:
            if m["role"] == "user":
                return "(DEMO)\n\n" + m["content"]
        return "(DEMO) Conteúdo indisponível."
    resp = client_oa.chat.completions.create(
        model=OPENAI_MODEL,
        messages=messages,
        temperature=0.2,
    )
    return resp.choices[0].message.content.strip()


# ============== DOCX FORMATTER ==============
def docx_formatar(texto: str, cliente: Optional[str], processo: Optional[str]) -> str:
    """Gera .docx simples e retorna apenas o nome do arquivo salvo em DOCS_DIR."""

    fn = f"peca_{uuid.uuid4().hex}.docx"
    path = DOCS_DIR / fn
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)

    # Cabeçalho institucional simples (pode trocar por template .dotx depois)
    doc.add_heading("MOURA MARTINS ADVOGADOS", 0)
    if cliente or processo:
        doc.add_paragraph(f"Cliente: {cliente or '—'} | Processo: {processo or '—'}")

    for par in texto.split("\n"):
        doc.add_paragraph(par)

    doc.save(path)
    return fn


# ============== PIPELINE PRINCIPAL ==============
def pipeline_gerar_peca(
    tipo_peca: str,
    facts: str,
    pedidos: str,
    documentos: List[str],
    cliente_id: Optional[str],
    processo_id: Optional[str],
    area: Optional[str] = None,
) -> Dict[str, Any]:
    """
    1) Busca similares no Chroma (modelos, peças, docs, juris internas)
    2) Busca jurisprudência (Tavily) filtrada por área e domínios de tribunais
    3) Monta prompt jurídico por área
    4) Chama GPT-5
    5) Gera .docx e retorna texto + link
    """

    # 1) RAG interno
    consulta = f"{tipo_peca}\n{facts}\n{pedidos}"
    similares = search_similares(consulta, k=4)

    # 2) Jurisprudência externa abrangente
    tav_q = f"{tipo_peca} {facts[:200]} {pedidos[:140]} jurisprudência aplicável"
    juris = buscar_tavily(tav_q, area=area, k=7)

    # 3) Prompt por área + contexto RAG
    ctx = RAGContext(
        tipo_peca=tipo_peca,
        facts=facts,
        pedidos=pedidos,
        documentos=documentos,
        cliente_id=cliente_id,
        processo_id=processo_id,
        area=area,
        similares=similares,
        jurisprudencia=juris,
    )
    messages = montar_messages(ctx)

    # 4) Geração
    texto = gerar_texto(messages)

    # 5) DOCX
    fname = docx_formatar(texto, cliente_id, processo_id)

    # 6) (Opcional) Aprendizado contínuo: indexa a peça recém-gerada
    try:
        upsert_peca(texto, cliente=cliente_id, processo=processo_id)
    except Exception:
        pass  # não quebra o fluxo se o upsert falhar

    return {"texto": texto, "docx": fname}


__all__ = [
    "DOCS_DIR",
    "pipeline_gerar_peca",
    "upsert_peca",
    "upsert_modelo",
    "upsert_juris",
    "upsert_doc_cliente",
]